"""
Sesión 6 de agentes · El agente que ficha papers

Primer entregable real: el agente lee un paper completo y produce una FICHA
comparativa estructurada (siempre los mismos campos), y la guarda en un
archivo Markdown junto al PDF.

Herramientas:
  - extraer_texto_pdf(ruta): todo el texto del PDF, con marcas de página.
  - guardar_ficha(ruta_salida, contenido): escribe la ficha en un archivo.

Requisitos:  uv add pypdf smolagents litellm  (ya instalados en sesiones previas)
"""

import os
import re
from pypdf import PdfReader
from docx import Document
from smolagents import tool, ToolCallingAgent, LiteLLMModel


def _ruta(ruta: str) -> str:
    return os.path.expanduser(ruta.strip())


# --- HERRAMIENTA 1: leer el PDF entero ---
@tool
def extraer_texto_pdf(ruta: str) -> str:
    """Extrae todo el texto de un PDF, marcando cada página con [Página N].

    Args:
        ruta: Ruta del PDF a leer.
    """
    try:
        lector = PdfReader(_ruta(ruta))
        partes = []
        for i, pagina in enumerate(lector.pages, start=1):
            texto = pagina.extract_text() or ""
            partes.append(f"[Página {i}]\n{texto}")
        completo = "\n\n".join(partes)
        # Límite generoso para no saturar; suficiente para un paper típico.
        return completo[:45000]
    except Exception as e:
        return f"Error al leer el PDF: {e}"


# --- HERRAMIENTA 2: guardar la ficha en Word (.docx) ---
def _md_a_docx(md: str, destino: str):
    """Convierte la ficha (en Markdown sencillo) a un documento Word.
    - '# Título'  -> encabezado
    - '- **Campo:** valor'  -> párrafo con el campo en negrita
    """
    doc = Document()
    for linea in md.splitlines():
        linea = linea.strip()
        if not linea:
            continue
        if linea.startswith("# "):
            doc.add_heading(linea[2:].strip(), level=1)
            continue
        # Quitamos el guion inicial de lista si lo hay.
        if linea.startswith("- "):
            linea = linea[2:].strip()
        p = doc.add_paragraph()
        # Partimos por los **...** para poner en negrita esos tramos.
        for i, trozo in enumerate(re.split(r"\*\*(.+?)\*\*", linea)):
            run = p.add_run(trozo)
            if i % 2 == 1:      # los tramos impares venían entre ** **
                run.bold = True
    doc.save(destino)


@tool
def guardar_ficha(ruta_salida: str, contenido: str) -> str:
    """Guarda la ficha como documento Word (.docx). Lo crea o lo sobrescribe.

    Args:
        ruta_salida: Ruta del .docx a escribir, por ejemplo '~/.../Ficha.docx'.
        contenido: El texto de la ficha en Markdown (con # y **negritas**).
    """
    try:
        destino = _ruta(ruta_salida)
        _md_a_docx(contenido, destino)
        return f"Ficha guardada en {destino}"
    except Exception as e:
        return f"Error al guardar: {e}"


model = LiteLLMModel(model_id="anthropic/claude-haiku-4-5")
agente = ToolCallingAgent(
    tools=[extraer_texto_pdf, guardar_ficha], model=model
)


# La plantilla de la ficha: SIEMPRE los mismos campos, para poder comparar.
PLANTILLA = """# Ficha: <título del artículo>

- **Referencia (APA):** autores (año). Título. Revista, volumen(número), páginas.
- **Objetivo:** qué pregunta de investigación responde.
- **Modelo / método:** qué modelo o técnica usa (SIR, EDP, calibración...).
- **Datos:** qué datos usa y de dónde.
- **Resultado principal:** el hallazgo clave, con cifras si las hay.
- **Limitaciones / hueco:** qué deja sin resolver.
- **Relevancia para mi tesis:** una frase sobre por qué me sirve.
"""


if __name__ == "__main__":
    pdf = "~/TESIS doctoral MAS/Publicaciones LEER/0 GRIPE_FLORIN/0 Gripe_Florin.pdf"
    salida = "~/TESIS doctoral MAS/Publicaciones LEER/0 GRIPE_FLORIN/Ficha - Gripe Florin.docx"

    tarea = (
        f"Lee el artículo científico en '{pdf}' usando extraer_texto_pdf. "
        "Luego redacta una ficha en español siguiendo EXACTAMENTE esta plantilla "
        f"(mismos campos, en Markdown):\n\n{PLANTILLA}\n\n"
        "Rellena cada campo basándote SOLO en el contenido del artículo, sin inventar. "
        "Mi tesis trata de modelizar la gripe con ecuaciones en derivadas parciales "
        "en el corredor Corea del Sur - Valencia, tenlo en cuenta para el campo de relevancia. "
        f"Finalmente, guarda la ficha con guardar_ficha en '{salida}'."
    )
    resultado = agente.run(tarea)
    print(f"\n==> {resultado}")
